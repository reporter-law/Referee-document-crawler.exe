"""程序说明"""
# -*-  coding: utf-8 -*-
# Author: cao wang
# Datetime : 2020
# software: PyCharm
# 收获:
import glob,os
from docx import Document
from win32com import client as wc
import tkinter as tk
from tkinter import *
import pdfplumber



def doc_to_docx(path):
    """
    目标函数
    """
    index_ = 1
    """打开wps"""
    """遍历doc文档"""
    for index,path_ in enumerate(glob.glob(os.path.join(path, "*.doc"))):

        try:
            word = wc.Dispatch('kwps.application')
            doc = word.Documents.Open(path_)
            # 错在没有加入文件名，位置加名字加格式
            """保存在同一文件夹下命名为docx，注意此处必须为docx否则转换不成功，必须为12否则无法读取"""
            doc.SaveAs2(path_.split('.')[0] + ".docx", 12)
            # print(path_.split('.')[0] + ".docx")
            # 保存在其他地方的 文件夹
            # doc.SaveAs(output + path_.split("/")[-1].split(".")[0] + ".docx", 16)
            message = ('正在进行格式转换--第%d页---------' % index)

            print("页面关闭第%d次---------------------------------------" % index_)
            index_ += 1
            doc.Close()
            yield message
            if index_ < 1000:
                pass
            elif index+1 == len(glob.glob(os.path.join(path, "*"))):
                print(len(glob.glob(os.path.join(path, "*"))))

            else:
                word.Quit()
                print("wps程序退出中--------------------%d---------" % index_)
                index_ = 1
        except Exception as e:
            print(e)

def read_pdf(input_path):
    """pdf转word"""
    input_paths = input_path.strip('\u202a')
    for path in glob.glob(os.path.join(input_paths,"*.pdf")):
        pdf = pdfplumber.open(path)
        text = []
        document = Document()
        for page in pdf.pages:
            '''pages为函数？'''
            document.add_paragraph(page.extract_text().strip().replace("\n", " "))  # 增加自然段add_paragraph
            document.add_page_break()  # 另起一页
        document.save(path+".docx")
        message = "docx转换成功-------"
        yield message

"""可视化"""
app = tk.Tk()  # 根窗口的实例(root窗口)
#调节窗口大小
app.geometry('450x450')
app.geometry('+450+150')
app.title('PDF格式转换工具')  # 根窗口标题
# label组件及文字内容 即具体设计
theLabel = tk.Label(app, text='格式转换工具！ ')
theLabel.pack()  # pack()用于自动调节组件的尺寸


"""输入框"""
theLabel = tk.Label(app, text='输入需要转换的路径！')
theLabel.pack()  # pack()用于自动调节组件的尺寸
entry=Entry(app,fg='blue',bg='white',width=50,bd=4)
entry.pack()
"""doc转docx"""
def run_doc_docx():
    path = entry.get().strip('\u202a')
    theLabel = tk.Label(app, text="转换路径为： "+path)
    theLabel.place(x=100,y=150) # pack()用于自动调节组件的尺寸
    message = doc_to_docx(path)
    for i in message:
        theLabel = tk.Label(app, text=i)
        theLabel.pack()  # pack()用于自动调节组件的尺寸

"""pdf转txt和docx"""
def run_pdf_txt():
    path = entry.get().strip('\u202a')
    theLabel = tk.Label(app, text="转换路径为： "+path)
    theLabel.place(x=150,y=130)
    message = read_pdf(path)
    for i in message:
        theLabel = tk.Label(app, text=i)
        theLabel.place(x=150,y=160)  # pack()用于自动调节组件的尺寸
button1 = Button(app, text="doc转docx", command=run_doc_docx,width=10,bg='white')
button1.place(x=100,y=90)
button3 = Button(app, text="pdf-docx", command=run_pdf_txt,bg='white')
button3.place(x=200,y=90)


app.mainloop()
app.quit()